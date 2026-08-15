"""
Gera o dicionário de dados do dataset IEEE-CIS Fraud Detection em .docx.
Cobre todas as colunas das duas tabelas: train_transaction e train_identity.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

SAIDA = Path(__file__).resolve().parent.parent / "reports" / "dicionario_dados.docx"

# ── Paleta de cores por grupo ──────────────────────────────────────────────────
CORES = {
    "alvo":        RGBColor(0x7C, 0x3A, 0xED),   # roxo
    "transacao":   RGBColor(0x05, 0x96, 0x69),   # verde
    "cartao":      RGBColor(0x0E, 0x7A, 0xFF),   # azul
    "endereco":    RGBColor(0xF5, 0x9E, 0x0B),   # âmbar
    "email":       RGBColor(0xEF, 0x44, 0x44),   # vermelho
    "C":           RGBColor(0x06, 0xB6, 0xD4),   # ciano
    "D":           RGBColor(0x84, 0xCC, 0x16),   # lima
    "M":           RGBColor(0xF9, 0x73, 0x16),   # laranja
    "V":           RGBColor(0x6B, 0x72, 0x80),   # cinza
    "identidade":  RGBColor(0xEC, 0x48, 0x99),   # rosa
    "dispositivo": RGBColor(0x14, 0xB8, 0xA6),   # teal
}

# ── Dados do dicionário ────────────────────────────────────────────────────────
GRUPOS = [

    {
        "titulo": "Variável Alvo",
        "arquivo": "train_transaction.csv",
        "cor": "alvo",
        "descricao": (
            "A única coluna que o modelo deve prever. "
            "Fortemente desbalanceada: ~96,5% são legítimas e ~3,5% são fraudes."
        ),
        "colunas": [
            ("isFraud", "int", "0 ou 1",
             "0 = transação legítima | 1 = fraude confirmada. "
             "Variável alvo do projeto. Nunca use como feature de entrada."),
        ],
    },

    {
        "titulo": "Identificadores e Tempo",
        "arquivo": "train_transaction.csv",
        "cor": "transacao",
        "descricao": (
            "Colunas de controle. TransactionDT não é um timestamp real — "
            "é um contador em segundos a partir de uma data de referência não divulgada pela Vesta."
        ),
        "colunas": [
            ("TransactionID",  "int",   "Ex: 2987000",
             "Chave primária. Usado para fazer o join com train_identity.csv."),
            ("TransactionDT",  "int",   "Ex: 86400 (= 1 dia)",
             "Segundos decorridos desde uma data-base secreta. "
             "Para extrair dia da semana e hora, use: dt % 86400 // 3600 = hora do dia."),
            ("TransactionAmt", "float", "Ex: 31.95, 117.00",
             "Valor da transação em dólares (USD). "
             "Fraudes tendem a ter valores específicos — analise a distribuição com escala log."),
            ("ProductCD",      "str",   "W, H, C, S, R",
             "Código do produto/serviço comprado. W é o mais frequente. "
             "Significado exato não divulgado, mas correlaciona com risco."),
        ],
    },

    {
        "titulo": "Informações do Cartão (card1–card6)",
        "arquivo": "train_transaction.csv",
        "cor": "cartao",
        "descricao": (
            "Dados sobre o cartão usado na transação. "
            "card4 e card6 têm valores legíveis; os demais são codificados pela Vesta."
        ),
        "colunas": [
            ("card1", "int",   "Ex: 4150, 9500",
             "Identificador do tipo de cartão (codificado). Alta cardinalidade — "
             "considere frequency encoding ou target encoding."),
            ("card2", "float", "Ex: 111, 360, 555",
             "Atributo adicional do cartão (codificado). Muitos valores nulos."),
            ("card3", "float", "Ex: 150, 185",
             "Atributo adicional do cartão (codificado)."),
            ("card4", "str",   "visa, mastercard, american express, discover",
             "Bandeira do cartão. Mais interpretável que os outros. "
             "Visa é a mais comum (~60% das transações)."),
            ("card5", "float", "Ex: 101, 145, 226",
             "Atributo adicional do cartão (codificado)."),
            ("card6", "str",   "debit, credit, debit or credit, charge card",
             "Tipo do cartão. Cartões de débito são mais frequentes."),
        ],
    },

    {
        "titulo": "Endereço (addr1, addr2)",
        "arquivo": "train_transaction.csv",
        "cor": "endereco",
        "descricao": (
            "Endereço associado à transação. "
            "Ambas as colunas são numéricas codificadas — não representam texto de endereço."
        ),
        "colunas": [
            ("addr1", "float", "Ex: 299, 325, 204",
             "Região de cobrança (billing region). Codificada numericamente. "
             "~13% de valores nulos."),
            ("addr2", "float", "Ex: 87, 96",
             "País de cobrança (billing country). "
             "Baixa variabilidade — maioria das transações é do mesmo país."),
            ("dist1", "float", "Ex: 0, 9, 298",
             "Distância entre o endereço de cobrança e o de entrega. "
             "Distâncias grandes podem indicar fraude. ~60% de nulos."),
            ("dist2", "float", "Ex: 0, 65",
             "Segunda medida de distância (entre outros endereços associados). "
             "~94% de nulos — use com cautela."),
        ],
    },

    {
        "titulo": "E-mail (P_emaildomain, R_emaildomain)",
        "arquivo": "train_transaction.csv",
        "cor": "email",
        "descricao": (
            "Domínios de e-mail do comprador (P = purchaser) e do destinatário (R = recipient). "
            "E-mails anônimos (gmail, yahoo) têm correlação diferente de e-mails corporativos."
        ),
        "colunas": [
            ("P_emaildomain", "str", "gmail.com, yahoo.com, hotmail.com…",
             "Domínio do e-mail de quem comprou. "
             "Domínios menos comuns ou de outros países → maior risco."),
            ("R_emaildomain", "str", "gmail.com, anonymous.com…",
             "Domínio do e-mail do destinatário. "
             "'anonymous.com' é especialmente relevante — alta taxa de fraude associada."),
        ],
    },

    {
        "titulo": "Grupo C — Variáveis de Contagem (C1–C14)",
        "arquivo": "train_transaction.csv",
        "cor": "C",
        "descricao": (
            "Contagens relacionadas ao histórico do cartão/endereço. "
            "Exemplos: quantos endereços diferentes já foram associados a esse cartão, "
            "quantos e-mails diferentes, etc. Significado exato não divulgado pela Vesta, "
            "mas inferido pela comunidade Kaggle."
        ),
        "colunas": [
            ("C1",  "float", "1–17",    "Nº de endereços associados ao cartão. C1=1 → cartão novo ou pouco usado."),
            ("C2",  "float", "1–17",    "Nº de endereços encontrados para o cartão (variante de C1)."),
            ("C3",  "float", "0–26",    "Contagem relacionada ao histórico do cartão (semântica não confirmada)."),
            ("C4",  "float", "0–15",    "Contagem de transações com o mesmo endereço de entrega."),
            ("C5",  "float", "0–5",     "Contagem de e-mails diferentes usados com esse cartão."),
            ("C6",  "float", "1–15",    "Contagem de endereços de cobrança distintos no histórico."),
            ("C7",  "float", "0–8",     "Contagem de transações recentes (janela de tempo curta)."),
            ("C8",  "float", "1–16",    "Contagem similar a C1/C2 — possivelmente por janela de tempo diferente."),
            ("C9",  "float", "0–7",     "Nº de transações com o mesmo endereço de entrega (variante)."),
            ("C10", "float", "0–14",    "Contagem de dispositivos diferentes usados com esse cartão."),
            ("C11", "float", "1–17",    "Contagem geral de histórico do cartão."),
            ("C12", "float", "0–15",    "Contagem de transações negadas anteriores."),
            ("C13", "float", "1–2918",  "Contagem de transações totais do comprador — alta cardinalidade."),
            ("C14", "float", "1–33",    "Contagem de endereços distintos (janela longa)."),
        ],
    },

    {
        "titulo": "Grupo D — Variáveis de Tempo (D1–D15)",
        "arquivo": "train_transaction.csv",
        "cor": "D",
        "descricao": (
            "Diferenças de tempo em dias entre eventos relacionados ao cartão/conta. "
            "D1 é a mais importante: representa há quantos dias o cartão foi visto pela última vez. "
            "Valores muito altos ou muito baixos podem indicar cartões recém-criados para fraude."
        ),
        "colunas": [
            ("D1",  "float", "0–640",  "Dias desde a transação anterior com esse cartão. "
                                       "⭐ Uma das features mais correlacionadas com isFraud."),
            ("D2",  "float", "0–640",  "Dias desde a transação anterior (variante — às vezes igual a D1)."),
            ("D3",  "float", "0–819",  "Dias desde que o endereço de entrega foi visto pela última vez."),
            ("D4",  "float", "0–869",  "Dias desde transação anterior (escopo diferente de D1)."),
            ("D5",  "float", "0–819",  "Dias desde o último uso do e-mail do comprador."),
            ("D6",  "float", "0–836",  "Variável de tempo (semântica incerta, ~88% nulos)."),
            ("D7",  "float", "0–843",  "Variável de tempo (~93% nulos)."),
            ("D8",  "float", "0–1708", "Dias desde o cadastro da conta (inferido pela comunidade)."),
            ("D9",  "float", "0–1",    "Fração do dia (hora/24) da última transação. Valor entre 0 e 1."),
            ("D10", "float", "0–876",  "Dias desde última transação no mesmo dispositivo."),
            ("D11", "float", "0–670",  "Dias desde o cadastro do comprador (~47% nulos)."),
            ("D12", "float", "0–649",  "Variável de tempo (~94% nulos)."),
            ("D13", "float", "0–820",  "Variável de tempo (~95% nulos)."),
            ("D14", "float", "0–855",  "Variável de tempo (~95% nulos)."),
            ("D15", "float", "0–879",  "Dias desde a última mudança de endereço (~44% nulos)."),
        ],
    },

    {
        "titulo": "Grupo M — Variáveis de Match (M1–M9)",
        "arquivo": "train_transaction.csv",
        "cor": "M",
        "descricao": (
            "Flags booleanas (True/False/NaN) que indicam se informações batem entre si. "
            "Por exemplo: o nome no cartão bate com o nome de cobrança? "
            "Divergências são sinais clássicos de fraude."
        ),
        "colunas": [
            ("M1", "str", "T, F, NaN", "Nome no cartão bate com nome de cobrança."),
            ("M2", "str", "T, F, NaN", "Atributo de match (semântica não confirmada)."),
            ("M3", "str", "T, F, NaN", "Endereço de cobrança bate com endereço no cadastro do cartão."),
            ("M4", "str", "M0–M6",     "Variável de match com múltiplos níveis (não apenas T/F)."),
            ("M5", "str", "T, F, NaN", "E-mail do comprador bate com e-mail cadastrado no cartão."),
            ("M6", "str", "T, F, NaN", "Match de outro atributo do cartão."),
            ("M7", "str", "T, F, NaN", "Match (~86% nulos)."),
            ("M8", "str", "T, F, NaN", "Match (~87% nulos)."),
            ("M9", "str", "T, F, NaN", "Match (~87% nulos)."),
        ],
    },

    {
        "titulo": "Grupo V — Features Vesta (V1–V339)",
        "arquivo": "train_transaction.csv",
        "cor": "V",
        "descricao": (
            "339 features proprietárias criadas pela Vesta Corporation. "
            "São combinações e transformações das outras colunas, "
            "mais sinais internos de risco que a Vesta não divulga. "
            "Muitas têm >90% de nulos. Agrupadas por faixas conforme análise da comunidade Kaggle."
        ),
        "colunas": [
            ("V1–V11",    "float", "0–1 ou contagens", "Relacionadas ao cartão e histórico de pagamento."),
            ("V12–V34",   "float", "0–1 ou contagens", "Relacionadas a endereço e localização."),
            ("V35–V52",   "float", "0–1 ou contagens", "Relacionadas ao dispositivo usado."),
            ("V53–V74",   "float", "0–1 ou contagens", "Relacionadas ao e-mail do comprador e destinatário."),
            ("V75–V94",   "float", "0–1 ou contagens", "Relacionadas ao valor da transação (TransactionAmt)."),
            ("V95–V137",  "float", "0–1 ou contagens", "Relacionadas a janelas de tempo e frequência."),
            ("V138–V166", "float", "0–1 ou contagens", "Features de contagem derivadas."),
            ("V167–V278", "float", "0–1 ou contagens", "Features mistas — maior concentração de nulos."),
            ("V279–V321", "float", "0–1 ou contagens", "Features mistas — correlação com identidade."),
            ("V322–V339", "float", "0–1 ou contagens", "Últimas features Vesta — menor variância, muitos nulos."),
        ],
    },

    {
        "titulo": "Identidade (id_01–id_38)",
        "arquivo": "train_identity.csv",
        "cor": "identidade",
        "descricao": (
            "Informações sobre o dispositivo, rede e comportamento do usuário. "
            "Presentes em apenas ~60% das transações (nem toda transação tem identity). "
            "Muitas colunas têm semântica parcialmente inferida pela comunidade."
        ),
        "colunas": [
            ("id_01", "float", "-5 a 0",            "Score de risco de identidade (negativo = mais seguro)."),
            ("id_02", "float", "0–526523",           "Score numérico de identidade (possível ID interno)."),
            ("id_03", "float", "0–824",              "Contagem relacionada à identidade."),
            ("id_04", "float", "0–429",              "Contagem relacionada à identidade."),
            ("id_05", "float", "0–480",              "Número de eventos de risco associados à conta."),
            ("id_06", "float", "-112 a 0",           "Score negativo — quanto menor, mais suspeito."),
            ("id_07", "float", "0–5",                "Contagem de correspondências de identidade."),
            ("id_08", "float", "0–5",                "Contagem de correspondências de identidade."),
            ("id_09", "float", "0–1",                "Flag de verificação de identidade."),
            ("id_10", "float", "0–1",                "Flag de verificação de identidade."),
            ("id_11", "float", "100–600",            "Score de conta/sessão."),
            ("id_12", "str",   "Found, NotFound",    "IP está em base de proxies conhecidos? Found = proxy."),
            ("id_13", "float", "1–2918",             "Contagem de transações associadas a esse dispositivo."),
            ("id_14", "float", "-480 a 480",         "Fuso horário do dispositivo (offset em minutos do UTC)."),
            ("id_15", "str",   "New, Found, Unknown","Status do browser/OS: New = recém-instalado."),
            ("id_16", "str",   "Found, NotFound",    "Autofill do browser usado? Found = sim."),
            ("id_17", "float", "100–600",            "Score de sessão/dispositivo."),
            ("id_18", "float", "0–1",                "Flag de identidade."),
            ("id_19", "float", "30–600",             "Score relacionado ao histórico do dispositivo."),
            ("id_20", "float", "30–600",             "Score relacionado ao histórico do dispositivo."),
            ("id_21", "float", "0–1",                "Flag (~97% nulos)."),
            ("id_22", "float", "0–1",                "Flag (~97% nulos)."),
            ("id_23", "str",   "TRANSPARENT",        "Tipo de proxy (~97% nulos)."),
            ("id_24", "float", "0–15",               "Score de conta (~97% nulos)."),
            ("id_25", "float", "1–2918",             "Contagem de uso do dispositivo."),
            ("id_26", "float", "1–2918",             "Contagem de uso do dispositivo."),
            ("id_27", "str",   "Found, NotFound",    "Verificação adicional de identidade."),
            ("id_28", "str",   "New, Found",         "Status de conta: New = conta recém-criada."),
            ("id_29", "str",   "Found, NotFound",    "Cookie habilitado no browser."),
            ("id_30", "str",   "Ex: Windows 10",     "Sistema operacional e versão do dispositivo."),
            ("id_31", "str",   "Ex: chrome 75.0",    "Browser e versão. Browsers desatualizados → maior risco."),
            ("id_32", "float", "Ex: 32",             "Profundidade de cor da tela (bits)."),
            ("id_33", "str",   "Ex: 1920x1080",      "Resolução da tela. Resoluções incomuns podem indicar bot."),
            ("id_34", "str",   "match_status:1",     "Status de match de identidade com valor numérico."),
            ("id_35", "str",   "T, F",               "Cookie de terceiros habilitado."),
            ("id_36", "str",   "T, F",               "Uso de proxy detectado (complemento de id_12)."),
            ("id_37", "str",   "T, F",               "Browser mobile."),
            ("id_38", "str",   "T, F",               "JavaScript habilitado no browser."),
        ],
    },

    {
        "titulo": "Dispositivo (DeviceType, DeviceInfo)",
        "arquivo": "train_identity.csv",
        "cor": "dispositivo",
        "descricao": (
            "Informações diretas sobre o dispositivo usado na transação."
        ),
        "colunas": [
            ("DeviceType", "str", "desktop, mobile",
             "Tipo do dispositivo. Mobile tem taxa de fraude levemente diferente de desktop."),
            ("DeviceInfo", "str", "Ex: Windows, iOS Device, Samsung…",
             "Modelo/sistema do dispositivo. Alta cardinalidade — agrupe em categorias "
             "(iOS, Android, Windows, Mac, Other) para usar como feature."),
        ],
    },
]


# ── Funções de formatação ──────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Define a cor de fundo de uma célula da tabela."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def rgb_to_hex(cor: RGBColor) -> str:
    return f"{cor[0]:02X}{cor[1]:02X}{cor[2]:02X}"


def cor_texto_para_fundo(cor: RGBColor) -> RGBColor:
    """Retorna branco ou preto conforme a luminosidade do fundo."""
    r, g, b = cor
    luminancia = (0.299 * r + 0.587 * g + 0.114 * b) / 255
    return RGBColor(0xFF, 0xFF, 0xFF) if luminancia < 0.5 else RGBColor(0x1E, 0x29, 0x3B)


def paragrafo(doc, texto, tamanho=11, negrito=False, italico=False,
              alinhamento=WD_ALIGN_PARAGRAPH.LEFT,
              cor=None, espaco_antes=0, espaco_depois=6):
    p = doc.add_paragraph(texto)
    p.alignment = alinhamento
    p.paragraph_format.space_before = Pt(espaco_antes)
    p.paragraph_format.space_after = Pt(espaco_depois)
    for run in p.runs:
        run.font.size = Pt(tamanho)
        run.font.bold = negrito
        run.font.italic = italico
        if cor:
            run.font.color.rgb = cor
    return p


def adicionar_grupo(doc, grupo):
    cor_obj = CORES[grupo["cor"]]
    cor_hex = rgb_to_hex(cor_obj)
    cor_txt = cor_texto_para_fundo(cor_obj)

    # Cabeçalho colorido do grupo
    p_titulo = doc.add_paragraph()
    p_titulo.paragraph_format.space_before = Pt(14)
    p_titulo.paragraph_format.space_after = Pt(4)
    run = p_titulo.add_run(f"  {grupo['titulo']}")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = cor_txt
    # Fundo colorido via XML
    pPr = p_titulo._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), cor_hex)
    pPr.append(shd)

    # Arquivo de origem
    p_arq = doc.add_paragraph()
    p_arq.paragraph_format.space_before = Pt(0)
    p_arq.paragraph_format.space_after = Pt(4)
    run2 = p_arq.add_run(f"  Arquivo: {grupo['arquivo']}")
    run2.font.size = Pt(9)
    run2.font.italic = True
    run2.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    # Descrição do grupo
    p_desc = doc.add_paragraph(grupo["descricao"])
    p_desc.paragraph_format.space_before = Pt(0)
    p_desc.paragraph_format.space_after = Pt(8)
    for run in p_desc.runs:
        run.font.size = Pt(10)
        run.font.italic = True

    # Tabela de colunas
    tabela = doc.add_table(rows=1, cols=4)
    tabela.style = 'Table Grid'

    # Cabeçalho da tabela
    cabecalhos = ["Coluna", "Tipo", "Valores exemplo", "Descrição"]
    for i, texto in enumerate(cabecalhos):
        cell = tabela.rows[0].cells[i]
        set_cell_bg(cell, cor_hex)
        p = cell.paragraphs[0]
        run = p.add_run(texto)
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = cor_txt
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)

    # Larguras das colunas
    larguras = [Cm(3.2), Cm(1.6), Cm(3.0), Cm(8.8)]
    for i, larg in enumerate(larguras):
        for row in tabela.rows:
            row.cells[i].width = larg

    # Linhas de dados
    for idx, (col, tipo, valores, descricao) in enumerate(grupo["colunas"]):
        row = tabela.add_row()
        dados = [col, tipo, valores, descricao]
        bg = "F8F9FC" if idx % 2 == 0 else "FFFFFF"
        for i, texto in enumerate(dados):
            cell = row.cells[i]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(texto)
            run.font.size = Pt(9)
            if i == 0:
                run.font.bold = True
                run.font.color.rgb = cor_obj
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph()  # espaço após tabela


def gerar_docx():
    doc = Document()

    # Margens
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    # ── Capa ──────────────────────────────────────────────────────────────────
    paragrafo(doc, "Dicionário de Dados", tamanho=20, negrito=True,
              alinhamento=WD_ALIGN_PARAGRAPH.CENTER,
              espaco_antes=0, espaco_depois=6)

    paragrafo(doc,
              "IEEE-CIS Fraud Detection Dataset",
              tamanho=13, italico=True,
              alinhamento=WD_ALIGN_PARAGRAPH.CENTER,
              cor=RGBColor(0x64, 0x74, 0x8B),
              espaco_antes=0, espaco_depois=4)

    paragrafo(doc,
              "TCC — Sistema Híbrido de Detecção de Fraudes com ML + RAG",
              tamanho=10,
              alinhamento=WD_ALIGN_PARAGRAPH.CENTER,
              cor=RGBColor(0x94, 0xA3, 0xB8),
              espaco_antes=0, espaco_depois=24)

    # ── Resumo do dataset ──────────────────────────────────────────────────────
    paragrafo(doc, "Visão Geral do Dataset", tamanho=13, negrito=True,
              espaco_antes=0, espaco_depois=6)

    resumo = (
        "O dataset IEEE-CIS Fraud Detection é composto por duas tabelas que devem ser unidas "
        "pela coluna TransactionID usando um left join (nem toda transação tem registro de identidade).\n\n"
        "• train_transaction.csv — 590.540 linhas × 394 colunas\n"
        "  Contém dados da transação em si: valor, cartão, endereço, e-mail e as features "
        "  dos grupos C, D, M e V.\n\n"
        "• train_identity.csv — 144.233 linhas × 41 colunas\n"
        "  Contém dados sobre o dispositivo e identidade digital do usuário. "
        "  Presente em ~60% das transações.\n\n"
        "• Taxa de fraude: 3,5% (20.663 fraudes de 590.540 transações)\n"
        "• Período coberto: 6 meses (inferido pelos valores de TransactionDT)\n"
        "• Fonte: Vesta Corporation — empresa de serviços de garantia de pagamento"
    )
    p = doc.add_paragraph(resumo)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(16)
    for run in p.runs:
        run.font.size = Pt(10)

    # ── Grupos de colunas ──────────────────────────────────────────────────────
    paragrafo(doc, "Descrição das Colunas por Grupo", tamanho=13, negrito=True,
              espaco_antes=0, espaco_depois=8)

    for grupo in GRUPOS:
        adicionar_grupo(doc, grupo)

    # ── Nota final ────────────────────────────────────────────────────────────
    doc.add_page_break()
    paragrafo(doc, "Notas Importantes para o TCC", tamanho=13, negrito=True,
              espaco_antes=0, espaco_depois=8)

    notas = [
        ("Colunas V1–V339 são anônimas",
         "A Vesta não divulgou o significado dessas features. Use análise de importância "
         "(SHAP) para identificar quais são relevantes. Não tente interpretar individualmente."),
        ("Muitos nulos não são erro",
         "Colunas como D6, D7, D12–D14 têm >90% de nulos porque nem toda transação gera "
         "esses eventos. O nulo em si é uma informação — crie uma flag is_null para essas colunas."),
        ("TransactionDT não é hora real",
         "Você pode extrair padrões de hora do dia (dt % 86400 // 3600) e dia da semana "
         "(dt // 86400 % 7), mas não sabe a data absoluta."),
        ("train_identity.csv só tem 60% das transações",
         "Após o left join, todas as colunas id_ e Device* terão NaN para os 40% sem identity. "
         "Trate esses NaNs com imputação ou indicadores de ausência."),
        ("card1–card5 têm alta cardinalidade",
         "Com milhares de valores únicos, one-hot encoding vai explodir a dimensionalidade. "
         "Prefira target encoding ou frequency encoding para essas colunas."),
    ]

    for titulo_nota, descricao_nota in notas:
        p_n = doc.add_paragraph()
        p_n.paragraph_format.space_before = Pt(0)
        p_n.paragraph_format.space_after = Pt(8)
        r1 = p_n.add_run(f"⚠ {titulo_nota}: ")
        r1.font.bold = True
        r1.font.size = Pt(10)
        r2 = p_n.add_run(descricao_nota)
        r2.font.size = Pt(10)

    SAIDA.parent.mkdir(parents=True, exist_ok=True)
    doc.save(SAIDA)
    print(f"Dicionário gerado: {SAIDA}")


if __name__ == "__main__":
    gerar_docx()
