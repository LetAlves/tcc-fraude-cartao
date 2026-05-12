"""
Gera o arquivo abstract_artigo.docx com o abstract do TCC nas versões
em português e inglês, formatado com estilos acadêmicos básicos.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

SAIDA = Path(__file__).resolve().parent.parent / "reports" / "abstract_artigo.docx"

ABSTRACT_PT = (
    "A detecção de fraudes em transações de cartão de crédito é um problema crítico "
    "no setor financeiro, caracterizado por severo desbalanceamento de classes e pela "
    "necessidade de explicações compreensíveis às partes interessadas. Este trabalho "
    "propõe um sistema híbrido que combina modelos de aprendizado de máquina com "
    "geração aumentada por recuperação (RAG) para detectar e explicar transações "
    "fraudulentas em linguagem natural. Utilizando o dataset IEEE-CIS Fraud Detection, "
    "com aproximadamente 590 mil transações e taxa de fraude de 3,5%, aplicamos "
    "técnicas de balanceamento via SMOTE e treinamos modelos XGBoost e LightGBM com "
    "ajuste de hiperparâmetros por Optuna. A interpretabilidade local das predições é "
    "obtida por meio de valores SHAP, cujas features mais relevantes são utilizadas "
    "como contexto para um pipeline RAG que recupera regras de domínio e gera "
    "explicações textuais por transação via modelo de linguagem. Os experimentos "
    "demonstram que o modelo [XGBoost/LightGBM] alcançou AUC-PR de [X,XX], "
    "F1-score de [X,XX] e Recall de [X,XX] a 5% de taxa de falsos positivos, "
    "superando os baselines tradicionais. As explicações geradas pelo componente RAG "
    "foram avaliadas quanto à fidelidade aos valores SHAP e à relevância de domínio, "
    "obtendo [métrica de qualidade]. O sistema proposto contribui para a área de IA "
    "explicável em finanças ao integrar, de forma inédita, modelos preditivos de alto "
    "desempenho com explicações contextualizadas em linguagem natural."
)

PALAVRAS_CHAVE_PT = (
    "Palavras-chave: detecção de fraude, cartão de crédito, aprendizado de máquina, "
    "geração aumentada por recuperação, explicabilidade, XGBoost, SHAP."
)

ABSTRACT_EN = (
    "Credit card fraud detection is a critical challenge in the financial sector, "
    "marked by severe class imbalance and the need for human-interpretable explanations. "
    "This paper proposes a hybrid system that integrates high-performance machine "
    "learning models with Retrieval-Augmented Generation (RAG) to detect and explain "
    "fraudulent transactions in natural language. Using the IEEE-CIS Fraud Detection "
    "dataset — approximately 590,000 transactions with a 3.5% fraud rate — we address "
    "class imbalance via SMOTE and train XGBoost and LightGBM classifiers with "
    "Optuna-based hyperparameter tuning. Local prediction interpretability is achieved "
    "through SHAP values, whose top features serve as context for a RAG pipeline that "
    "retrieves domain-specific rules and generates per-transaction textual explanations "
    "via a large language model. Experimental results show that the [XGBoost/LightGBM] "
    "model achieves an AUC-PR of [X.XX], F1-score of [X.XX], and Recall of [X.XX] at "
    "a 5% false positive rate, outperforming traditional baselines. RAG-generated "
    "explanations are evaluated on SHAP fidelity and domain relevance, yielding "
    "[quality metric]. The proposed system advances explainable AI in finance by "
    "uniquely combining predictive accuracy with contextualised, human-readable "
    "fraud explanations."
)

PALAVRAS_CHAVE_EN = (
    "Keywords: fraud detection, credit card, machine learning, "
    "retrieval-augmented generation, explainability, XGBoost, SHAP."
)

NOTA = (
    "Nota: os valores entre colchetes [ ] devem ser preenchidos após a conclusão "
    "dos experimentos (previsto para julho/2026)."
)


def adicionar_linha_horizontal(doc):
    """Insere uma linha horizontal (borda inferior) entre seções."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def configurar_paragrafo(paragrafo, tamanho=11, negrito=False, italico=False,
                          alinhamento=WD_ALIGN_PARAGRAPH.JUSTIFY, cor=None,
                          espaco_antes=0, espaco_depois=6):
    paragrafo.alignment = alinhamento
    paragrafo.paragraph_format.space_before = Pt(espaco_antes)
    paragrafo.paragraph_format.space_after = Pt(espaco_depois)
    for run in paragrafo.runs:
        run.font.size = Pt(tamanho)
        run.font.bold = negrito
        run.font.italic = italico
        if cor:
            run.font.color.rgb = RGBColor(*cor)


def gerar_docx():
    doc = Document()

    # Margens acadêmicas (ABNT: 3cm esq/sup, 2cm dir/inf)
    for section in doc.sections:
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    # Fonte padrão do documento
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(12)

    # ── Título ────────────────────────────────────────────────────────────────
    titulo = doc.add_paragraph(
        "Sistema Híbrido de Detecção e Explicação de Fraudes em Transações "
        "de Cartão de Crédito Combinando Machine Learning e RAG"
    )
    configurar_paragrafo(titulo, tamanho=14, negrito=True,
                         alinhamento=WD_ALIGN_PARAGRAPH.CENTER,
                         espaco_antes=0, espaco_depois=4)

    # Autores
    autores = doc.add_paragraph("Letícia Alves Carvalho · Lucas [Sobrenome]")
    configurar_paragrafo(autores, tamanho=11, italico=True,
                         alinhamento=WD_ALIGN_PARAGRAPH.CENTER,
                         espaco_antes=0, espaco_depois=2)

    # Instituição
    inst = doc.add_paragraph("[Nome da Universidade] — Curso de Ciência da Computação")
    configurar_paragrafo(inst, tamanho=10,
                         alinhamento=WD_ALIGN_PARAGRAPH.CENTER,
                         cor=(100, 100, 100), espaco_antes=0, espaco_depois=16)

    adicionar_linha_horizontal(doc)

    # ── Resumo (PT) ───────────────────────────────────────────────────────────
    rotulo_pt = doc.add_paragraph("Resumo")
    configurar_paragrafo(rotulo_pt, tamanho=11, negrito=True,
                         alinhamento=WD_ALIGN_PARAGRAPH.LEFT,
                         espaco_antes=12, espaco_depois=4)

    corpo_pt = doc.add_paragraph(ABSTRACT_PT)
    configurar_paragrafo(corpo_pt, tamanho=11, espaco_antes=0, espaco_depois=6)
    # Recuo de 1,25 cm nas duas margens (estilo ABNT para resumo)
    corpo_pt.paragraph_format.left_indent = Cm(1.25)
    corpo_pt.paragraph_format.right_indent = Cm(1.25)

    pc_pt = doc.add_paragraph(PALAVRAS_CHAVE_PT)
    configurar_paragrafo(pc_pt, tamanho=10, italico=True,
                         espaco_antes=0, espaco_depois=16)
    pc_pt.paragraph_format.left_indent = Cm(1.25)
    pc_pt.paragraph_format.right_indent = Cm(1.25)

    adicionar_linha_horizontal(doc)

    # ── Abstract (EN) ─────────────────────────────────────────────────────────
    rotulo_en = doc.add_paragraph("Abstract")
    configurar_paragrafo(rotulo_en, tamanho=11, negrito=True,
                         alinhamento=WD_ALIGN_PARAGRAPH.LEFT,
                         espaco_antes=12, espaco_depois=4)

    corpo_en = doc.add_paragraph(ABSTRACT_EN)
    configurar_paragrafo(corpo_en, tamanho=11, espaco_antes=0, espaco_depois=6)
    corpo_en.paragraph_format.left_indent = Cm(1.25)
    corpo_en.paragraph_format.right_indent = Cm(1.25)

    pc_en = doc.add_paragraph(PALAVRAS_CHAVE_EN)
    configurar_paragrafo(pc_en, tamanho=10, italico=True,
                         espaco_antes=0, espaco_depois=16)
    pc_en.paragraph_format.left_indent = Cm(1.25)
    pc_en.paragraph_format.right_indent = Cm(1.25)

    adicionar_linha_horizontal(doc)

    # ── Nota sobre placeholders ───────────────────────────────────────────────
    nota_p = doc.add_paragraph(f"⚠ {NOTA}")
    configurar_paragrafo(nota_p, tamanho=9, italico=True,
                         alinhamento=WD_ALIGN_PARAGRAPH.LEFT,
                         cor=(150, 80, 0), espaco_antes=12, espaco_depois=0)

    # Salva
    SAIDA.parent.mkdir(parents=True, exist_ok=True)
    doc.save(SAIDA)
    print(f"Arquivo gerado: {SAIDA}")


if __name__ == "__main__":
    gerar_docx()
